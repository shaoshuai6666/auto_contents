from translation import translate_paragraph
from utils import insert_paragraph_after, delete_paragraph
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Cm

# 删除Contents和下一个一级标题之间的内容

def delete_between_contents_and_heading1(paragraphs):
    found_contents = False
    for i,paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if text.startswith("Contents"):
            contents_index = i
            found_contents = True
        elif found_contents and paragraph.style.name.startswith("Heading 1"):
            insert_paragraph_after(paragraphs[contents_index], text=" ", style=None,tab_stops_position=None)
            break
        elif found_contents:
            delete_paragraph(paragraph)
    return contents_index  

def translate_contents(doc,translation_service):
    
    contents_index = delete_between_contents_and_heading1(doc.paragraphs)

    paragraphs = doc.paragraphs  # 获取所有段落


    # 初始化标志变量
    is_contents = False
    found_mulu = False
    #翻译
    for i, para in enumerate(paragraphs):
        # 找到目录
        if para.text.find("目　　录") != -1:
            found_mulu = True
            continue
        # 判断段落样式
        elif found_mulu and para.style.name.startswith("toc"):
            print(para.text)
            result = translate_paragraph(para,translation_service)

            #把页码前面的空格替换为tab
            last_space_index = result.rfind(' ')# 查找最右边的空格
            if last_space_index >= 0:
                result = result[:last_space_index] + '\t' + result[last_space_index+1:]# 将最右边的空格替换为tab制表符
                
            print(result)
            Contents_style = doc.styles[para.style.name]
            paragraphs[contents_index+1].insert_paragraph_before(result, Contents_style)

        elif found_mulu and para.text.strip().startswith("Contents"):
            
            found_mulu = False
            break
        
    ###一般在中文目录中既有"摘要"又有"Abstract"，而英文目录中只需要一个"Abstract"，所以把重复翻译的摘要目录删掉
    paragraphs = doc.paragraphs  # 重新获取所有段落
    if "Abstract" in paragraphs[contents_index+1].text and "Abstract" in paragraphs[contents_index+2].text:# 如果Contents中第一行和第二行同时存在“Abstract”
        delete_paragraph(paragraphs[contents_index+1])#则删除第一行由"摘要"翻译的"Abstract"

    ###添加制表符
    paragraphs = doc.paragraphs  # 重新获取所有段落
    for i, para in enumerate(paragraphs):
        # 找到目录
        # print(i,is_contents,para.text)
        if para.text.strip().startswith("Contents"):
            is_contents = True
            # print("is_contents = ",is_contents)
            continue
        # 判断段落样式
        elif is_contents and not para.style.name.startswith("toc"):
            is_contents = False
            break
        elif is_contents and para.style.name.startswith("toc"):
            # print(para.text)
            para.paragraph_format.tab_stops.add_tab_stop(Cm(14.99), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS)  # 右对齐，前导字符为点
            
        
        